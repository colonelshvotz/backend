from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI

from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt
import uuid
import os

import json
from typing import Optional

from fastapi import HTTPException

from fastapi import Request

from fastapi import Body

import random

import uvicorn

import base64

client = OpenAI(api_key=os.environ['StoryT'])

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://frontend-rho-snowy-74.vercel.app", "http://localhost:3000"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def root():
    return {
        "message":
        "Welcome to the FastAPI backend!",
        "endpoints": [
            "/load-book",
            "/start",
            "/continue",
            "/generate-moment-image",
            "/export-story",
            "/test-chapter",
            "/developer/trait-scores",
            "/developer/empathy-scores",
            "/save-progress",
        ]
    }


# Simple session storage
session = {
    "genre": None,
    "character_name": None,
    "character_description": None,
    "history": [],
    "image_url": None
}


# Request schemas
class StartRequest(BaseModel):
    genre: str
    character_name: str
    character_description: str
    story_idea: str = ""
    let_ai_decide: bool = False
    skip_image: bool = False
    narrator_style: Optional[str] = "neutral"  # new


class ContinueRequest(BaseModel):
    user_input: str


# Load books
class BookLoadRequest(BaseModel):
    title: str
    genre: Optional[str] = None
    character_name: str
    character_description: Optional[str] = None
    story_idea: Optional[str] = None
    let_ai_decide: Optional[bool] = False
    skip_image: Optional[bool] = False


# Books and worlds created and saved
BOOKS_DIR = "books"
os.makedirs(BOOKS_DIR, exist_ok=True)


def get_book_path(title: str):
    safe_title = title.replace(" ", "_")
    return os.path.join(BOOKS_DIR, f"{safe_title}.json")


def load_book_data(title: str) -> dict:
    path = get_book_path(title)
    if os.path.exists(path):
        with open(path, "r") as f:
            book = json.load(f)

        if "chapters" not in book or not isinstance(book["chapters"], list):
            book["chapters"] = []
        return book
    return {}


def save_book_data(data: dict):
    path = get_book_path(data["title"])
    with open(path, "w") as f:
        json.dump(data, f, indent=2)


# Chapters and summaries inside books
def create_chapter_summary_and_story(character_name: str, genre: str,
                                     history: list):
    if not history:
        raise ValueError("No story history found for this character.")

    # Filter out /game and user inputs
    #story_lines = [
       # line for line in history if not line.strip().startswith(">")
       # and not line.strip().lower().startswith("/game")
    #]
    #full_log = "\n".join(story_lines)
    full_log = history

    # Ask GPT to summarize in 1 sentence
    summary_prompt = (
        f"Summarize the following story in 1 sentence. The genre is {genre}. "
        f"The main character is {character_name}. Focus on the arc of the plot:\n\n{full_log}"
    )

    summary_response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": summary_prompt
        }],
        temperature=0.7,
    )

    summary_text = summary_response.choices[0].message.content.strip(
    ) if summary_response.choices[0].message.content else "No summary available"

    # Ask GPT to rewrite the story in third-person past tense
    rewrite_prompt = (
        f"Rewrite the following story log as a third-person limited past-tense narrative. "
        f"The protagonist is named {character_name}. Genre: {genre}. "
        #f"Only include what {character_name} could reasonably know or perceive. "
        f"Describe other characters’ actions and dialogue in vivid, cinematic prose. You are an expert, literary author."
        f"Remove all second-person references, user prompts, and transform it into polished storybook prose. "
        f"Here is the story log:\n\n{full_log}")

    prose_response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": rewrite_prompt
        }],
        temperature=0.7,
    )

    prose_text = prose_response.choices[0].message.content.strip(
    ) if prose_response.choices[
        0].message.content else "Failed to generate rewritten story."

    return {"summary": summary_text, "prose": prose_text}


def save_chapter_to_book(book_title: str, character_name: str, summary: str,
                         story_text: str):

    book = session.get("book")
    character_name = session.get("character_name", "Unnamed Hero")
    character_description = session.get("character_description", "")
    genre = session.get("genre", "Unknown")
    history = session.get("history", [])
    book_title = session.get("book", {}).get("title")

    #Test print lines
    print("Saving chapter to book...")
    print("Book:", book_title)
    print("Character:", character_name)
    print("Summary:", summary[:100])

    if not book or not history:
        return

    # Filter out /game commands
    #filtered_history = []
    #skip_next = False
    #for line in history:
       # if line.strip().lower().startswith("/game") or skip_next:
        #    skip_next = not skip_next
        #    continue
       # filtered_history.append(line)

   # full_log = "\n".join(filtered_history)

    # Get third-person version
    # Removed filtered history, just using full history
    prompt_rewrite = (
        f"Rewrite the following story log as a third-person limited past-tense narrative. "
        f"The protagonist is named {character_name}. Genre: {genre}. "
        f"Only include what {character_name} could reasonably know or perceive. "
        f"Describe other characters’ actions and dialogue in vivid, cinematic prose. "
        f"Remove all second-person references, user prompts, and transform it into polished storybook prose. Write like an expert, published author."
        f"Here is the story log:\n\n{history}")

    response_rewrite = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt_rewrite
        }],
        temperature=0.7,
    )

    story_text = response_rewrite.choices[0].message.content.strip(
    ) if response_rewrite.choices[
        0].message.content else "Failed to generate rewritten story."

    # Get summary
    prompt_summary = (
        f"Summarize the following story in exactly one sentence:\n\n{prompt_rewrite}"
    )

    response_summary = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt_summary
        }],
        temperature=0.5,
    )

    if response_summary.choices[0].message.content:
        summary = response_summary.choices[0].message.content.strip()
    else:
        summary = "No summary available"

    # Prepare chapter data
    chapter_data = {
        "character": character_name,
        "summary": summary,
        "story": story_text
    }

    # Save to book
    if "chapters" not in book or not isinstance(book["chapters"], list):
        book["chapters"] = []

    book["chapters"].append(chapter_data)
    save_book_data(book)


# Optional tool to clean old book files and store characters as a list if they've been stored as an array
def sanitize_books():
    for filename in os.listdir(BOOKS_DIR):
        if filename.endswith(".json"):
            path = os.path.join(BOOKS_DIR, filename)
            with open(path, "r") as f:
                book = json.load(f)

            updated = False
            if not isinstance(book.get("characters"), list):
                book["characters"] = []
                updated = True
            if not isinstance(book.get("stories"), list):
                book["stories"] = []
                updated = True

            if updated:
                with open(path, "w") as f:
                    json.dump(book, f, indent=2)


# Generate character image from full context
def generate_character_image(prompt):
    try:
        response = client.images.generate(model="dall-e-3",
                                          prompt=prompt,
                                          size="1024x1024",
                                          n=1)
        return response.data[0].url
    except Exception as e:
        print("Image generation failed:", e)
        return None


# Persistent character data between sessions
CHARACTERS_DIR = "characters"
os.makedirs(CHARACTERS_DIR, exist_ok=True)


def get_character_path(book_title: str, character_name: str):
    safe_book = book_title.replace(" ", "_")
    safe_name = character_name.replace(" ", "_")
    return os.path.join(CHARACTERS_DIR, f"{safe_book}__{safe_name}.json")


def load_character(book_title: str, character_name: str) -> dict:
    path = get_character_path(book_title, character_name)
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return {}


def save_character(book_title: str, character_name: str, character_data: dict):
    path = get_character_path(book_title, character_name)
    with open(path, "w") as f:
        json.dump(character_data, f, indent=2)


#NPC system to remember user characters and their interactions
NPC_MEMORY_DIR = "npc_memory"
os.makedirs(NPC_MEMORY_DIR, exist_ok=True)


def get_npc_memory_path(book_title: str, character_name: str):
    safe_book = book_title.replace(" ", "_")
    safe_char = character_name.replace(" ", "_")
    return os.path.join(NPC_MEMORY_DIR,
                        f"npc_memory__{safe_book}__{safe_char}.json")


def load_npc_memory(book_title: str, character_name: str) -> dict:
    path = get_npc_memory_path(book_title, character_name)
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return {}


def save_npc_memory(book_title: str, character_name: str, memory_data: dict):
    path = get_npc_memory_path(book_title, character_name)
    with open(path, "w") as f:
        json.dump(memory_data, f, indent=2)


def update_npc_memory_from_story(book_title: str, character_name: str,
                                 latest_story: str):
    prompt = (
        "You are an AI system that tracks memories of recurring characters (NPCs) in a story. "
        "Given the following scene, identify any important characters who are NOT the player. "
        "For each one, extract:\n"
        "1. Name\n2. Their role or personality\n3. Any important memory they now have of the player.\n\n"
        f"Scene:\n{latest_story}\n\n"
        "Respond in JSON format like this:\n"
        "{\n"
        "  \"Sheriff Harlan\": {\n"
        "    \"role\": \"Grizzled lawman of Red Hollow\",\n"
        "    \"memory\": \"The player saved him during a bank robbery. He owes them a favor.\"\n"
        "  },\n"
        "  ... more NPCs\n"
        "}")

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.6,
    )

    try:
        if response.choices[0].message.content:
            npc_update = json.loads(response.choices[0].message.content)
        else:
            npc_update = {}
    except Exception:
        npc_update = {}

    memory_data = load_npc_memory(book_title, character_name)

    for npc, details in npc_update.items():
        memory_data[npc] = details  # Overwrite with latest memory

    save_npc_memory(book_title, character_name, memory_data)


# --- Mystery/Investigation Mode ---

MYSTERY_DIR = "mysteries"
os.makedirs(MYSTERY_DIR, exist_ok=True)


def get_mystery_path(book_title: str, character_name: str):
    safe_book = book_title.replace(" ", "_")
    safe_character = character_name.replace(" ", "_")
    return os.path.join(MYSTERY_DIR,
                        f"mystery_{safe_book}__{safe_character}.json")


def save_mystery_state(book_title: str, character_name: str,
                       mystery_data: dict):
    path = get_mystery_path(book_title, character_name)
    with open(path, "w") as f:
        json.dump(mystery_data, f, indent=2)


def load_mystery_state(book_title: str, character_name: str):
    path = get_mystery_path(book_title, character_name)
    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    return None


def generate_suspects(genre: str) -> list:
    prompt = (
        f"Generate a list of 3 suspects for a mystery story set in a {genre} world. "
        f"Each suspect should include a name, their relationship to the victim, a motive, and an alibi. "
        f"Format it in JSON like this:\n"
        "[\n"
        "  {\n"
        "    \"name\": \"Suspect Name\",\n"
        "    \"relation\": \"Their relationship to the victim\",\n"
        "    \"motive\": \"Their motive\",\n"
        "    \"alibi\": \"Their alibi\"\n"
        "  },\n"
        "  ...more suspects\n"
        "]")

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.7,
    )

    if response.choices[0].message.content:
        suspects = json.loads(response.choices[0].message.content.strip())
        return suspects
    else:
        print("Error parsing suspects: No content returned from OpenAI.")
        return []


def generate_initial_clue(suspects: list) -> str:
    suspect_names = [s["name"] for s in suspects]
    prompt = (
        f"Create the first clue in a mystery involving the following suspects: {', '.join(suspect_names)}. "
        f"The clue should be vague but meaningful, and it should implicate at least one suspect, "
        f"but not reveal the culprit. Describe it in 2–3 sentences, written in immersive, cinematic language."
    )

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.85,
    )

    return response.choices[0].message.content.strip() if response.choices[
        0].message.content else "Failed to generate rewritten story."


#Create a list of suspects that will be in the mystery that you can then accuse
def extract_suspects_from_story(story_text: str) -> list:
    prompt = (
        "You are an AI mystery assistant. Analyze the following opening passage of a mystery story and extract 3 to 5 potential suspects. "
        "These should be characters who were mentioned by name and could plausibly be involved in the central mystery.\n\n"
        f"Story:\n{story_text}\n\n"
        "Return ONLY a JSON list of full names like this:\n[\"Harper Vane\", \"Dr. Morley\", \"Cassius Black\"]"
    )

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.6,
    )

    if response.choices[0].message.content:
        suspects = json.loads(response.choices[0].message.content)
        return suspects
    else:
        print("Error parsing suspects: No content returned from OpenAI.")
        return []


#Inventory concept - doesn't really work because the story engine lets you have any item you want
#def generate_starting_inventory(genre: str) -> list:
#    prompt = (
#        f"Generate a short list of 3 unique starting inventory items #appropriate for a character "
#        f"in a {genre} setting. Keep the list compact, with only the item names. "
#        f"Do not include descriptions or numbers.")

#    response = client.chat.completions.create(
#        model="gpt-3.5-turbo",
#        messages=[{
#            "role": "user",
#            "content": prompt
#        }],
#        temperature=0.7,
#    )

#    text = response.choices[0].message.content.strip(
#    ) if response.choices[0].message.content is not None else ""
#    items = [
#        item.strip("-• ").strip() for item in text.splitlines()
#        if item.strip()
#    ]
#    return items


def create_or_load_character(book_title: str, character_name: str, genre: str,
                             description: str):
    existing = load_character(book_title, character_name)
    if existing:
        return existing
    else:
        new_character = {
            "name": character_name,
            "description": description,
            #"inventory": generate_starting_inventory(genre),
            "reputation": {},
            "relationships": {},
            "empathy_scores": [],
            "Empathy": [],
            "Risk-taking": [],
            "Violence tendancy": [],
            "Loyalty": [],
            "Curiosity": [],
            "Rule-breaking": [],
            "Sacrifice": [],
            "Openness": [],
            "Conscientiousness": [],
            "Extroversion": [],
            "Agreeableness": [],
            "Neuroticism": [],
        }
        save_character(book_title, character_name, new_character)
        return new_character


# Generate story segments
def generate_story(genre, user_input=None, character_name=None):
    character_name = character_name or session.get("character_name",
                                                   "Your character")
    book_title = session.get("book", {}).get("title", "")

    # Retrieve current plot step from skeleton
    plot_step = ""
    skeleton = session.get("plot_skeleton", [])
    step = session.get("current_step", 0)

    if 0 <= step < len(skeleton):
        plot_step = skeleton[step]

    #NPC loads their relationship to user's characters
    npc_memories = load_npc_memory(book_title, character_name)
    npc_memory_text = "\n".join([
        f"{name} — {data['role']}. Memory: {data['memory']}"
        for name, data in npc_memories.items()
    ])

    if user_input is None:
        story_idea = session.get("story_idea", "")
        let_ai_decide = session.get("let_ai_decide", False)

        if let_ai_decide or not story_idea.strip():
            idea_part = (
                "The AI should invent the plot and location entirely, using the genre and character. "
                "Internally consider 10 different plots and settings, choose one at random, and begin the story as if this is the world the player awakens into. "
                "Do not show the list of plots or locations to the user. Let the story unfold naturally."
            )  #add plot skeleton
            if plot_step:
                idea_part += f" Begin with this initial story beat: {plot_step}"

        else:
            idea_part = f"The story should be about: {story_idea.strip()} and go in a direction unique to any previous history of the story or character"
            #add plot skeleton
            if plot_step:
                idea_part += f" Begin with this initial story beat: {plot_step}"

        # Fetch all past chapters for the current book, regardless of character
        # Fetch only past chapters for the *current* character
        past_chapters = [
            ch["story"] for ch in session.get("book", {}).get("chapters", [])
            if ch.get("character") == character_name
        ]

        if past_chapters:
            recap_text = "\n---\n".join(past_chapters)
            previously = (
                f"\n\nThe following events occurred in past stories within this world. "
                f"Do not repeat these plots, but allow characters, consequences, and themes to return or evolve:\n{recap_text}\n"
            )
        else:
            previously = ""

        if genre.lower() == "mystery":
            prompt = (
                f"You are beginning an interactive detective mystery in the style of Agatha Christie. "
                f"The main character is {character_name}, a clever and observant sleuth. "
                f"Write in second-person, but keep the prose elegant, sharp, and focused on deduction and subtle tension. "
                f"The setting should have intrigue — perhaps a small town, an old mansion, or a quiet train. "
                f"Introduce a crime, a few suspicious characters, and build toward a trail of clues and red herrings. "
                f"{idea_part} End with a question or moment of uncertainty, driving curiosity and analysis."
                f"{previously}")
        else:

            prompt = (
                f"You are beginning an immersive, second-person interactive story in the genre '{genre}'. "
                f"The main character’s name is {character_name}, but narration should use 'you' and stay in second-person. "
                f"{idea_part} Start with cinematic sensory detail. End with a dramatic decision or question."
                f"Begin with a cinematic opening — a scene full of tension, color, and implied conflict. Introduce at least one thing that seems off, suspicious, or dangerous. "
                f"Include micro details that make the world specific and memorable. End with a dramatic choice or uncomfortable question."
                f"{previously}"  # <-- added here to help the AI avoid past storylines
            )

    else:
        history = "\n".join(session["history"][-5:])
        prompt = (
            f"This is an ongoing {genre} story. The character's name is {character_name}. "
            f"Use immersive second-person narration. Previous context:\n{history}\n"
            f"Current story goal:\n{plot_step}\n\n"  #plot skeleton add
            f"The player says: '{user_input}'. DO NOT start the story with this idea, but gradually build to it.  Continue the story and end with a prompt."
        )

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role":
            "system",
            "content":
            ("You are an expert storyteller guiding the player through an immersive, interactive narrative. "
             "Your role is to narrate vivid, emotionally rich, second-person prose. The player experiences the world directly — never refer to 'the player'; simply describe what 'you' see, hear, feel, or do. "
             "Your storytelling should feel cinematic and visceral. Use micro worldbuilding to describe textures, sounds, smells, and background characters. Let the world feel lived-in, not generic. Every paragraph should include at least one fun, unexpected detail — a strange character quirk, an overheard rumor, or an eerie object in the room."
             "Build tension constantly. Give every scene a question or danger — something that leaves the player slightly uneasy. When conflict arises, escalate it with consequences. Minor victories should reveal deeper mysteries or hidden costs. Character dialogue should hint at hidden motives, lies, or vulnerabilities."
             "The story should follow principles of professional screenwriting and narrative design:\n"
             "- Stories are about conflict and the human desire to resolve it\n"
             "- Conflict can be generated when characters do the wrong thing for the right reason\n"
             "- Structure is based on the three-act format, but the player may deviate organically. The character starts in stasis, their lives have been unchanged for years leading up to this first moment.  Then they experience an event that will call them to action, something that promises an adventure.  Then the actions of the story drive the user to either a big win or a big loss.  Then the bad guys or bad element is hot on thier tail, which drives them to another big loss.  Then they see a final big bad boss that they'll have to defeat.  They have to mount a final offensive to slip past defenses or fight past defenses.  Then defeat the final big bad boss, thus ending the story.\n"
             "- Early choices should promise adventure; later choices should drive toward resolution or disaster\n"
             "- The hero’s hope must always present a better path forward, even when unclear\n"
             "- Give the player minor victories and setbacks to shape emotional arcs\n"
             "- Rarely, a bad decision may result in the player’s death — if so, allow them to restart or rewind\n"
             "- Do not break the fourth wall.\n"
             "- Avoid any sexual content. If the player attempts to introduce it, steer away naturally without drawing attention.\n"
             "- Mature content such as violence or death is allowed, but respect tone and genre.\n"
             "\n"
             "DO NOT gloss over story moments such as 'the journey was long with unexpected twists and turns'.  Turn that journey into an adventure for the user and delve into it even if it diverts from the main story.\n"
             "DO NOT be vague with introducing characters or villains.  Be specific about descriptions of clothing and what type of creature it is: human, mythological like elf or giant, or anything else."
             "Randomly make the user fail at something they attempt to do.  Failure lets the user improvise new solutions."
             "Roughly 1 in 5 items and/or offerings NPCs give to the user's character should be traps or at least affect the user negatively."
             "The story world is persistent. Characters, places, and events from earlier moments may return later, changed by time. If a user returns to this world later, they may start in the future.\n"
             "NPCs have motivations, goals, and memory. They may betray the player if their goals diverge, but the player may find clues to prevent betrayal.\n"
             "Some decisions should be open-ended (e.g., 'Where do you want to search?' or 'What do you want to say?').\n"
             "At every scene's end, ask the player what they want to do next."
             f"\n\nNPCs in this world remember past interactions. Here are their current memories:\n{npc_memory_text}"
             "If the story is approaching the final act, you must prioritize narrative closure. Begin resolving major plotlines, allow the character to succeed or fail meaningfully, and guide the story to a full, emotionally satisfying ending.\n"
             "Once the story naturally concludes, write a final sentence ending the narrative and add 'End of Chapter.'\n"
             "Do not invent new conflicts once the story is ready to end."
             )
        }, {
            "role": "user",
            "content": prompt
        }],
        temperature=0.9,
        max_tokens=500)

    story = response.choices[0].message.content.strip(
    ) if response.choices[0].message.content else "No story generated."
    return story


def handle_game_command(command: str) -> str:
    command = command.lower()

    if command == "end":
        genre = session.get("genre")
        character_name = session.get("character_name")
        book_title = session.get("book", {}).get("title")

        if not (genre and character_name and book_title):
            return "Unable to save chapter — missing session data."

        # Only add "The End." once
        if not session["history"] or session["history"][-1].strip(
        ) != "The End.":
            session["history"].append("The End.")

        try:
            chapter_info = create_chapter_summary_and_story(
                character_name, genre, session["history"])
            save_chapter_to_book(book_title, character_name,
                                 chapter_info["summary"],
                                 chapter_info["prose"])
            return "The End."
        except Exception as e:
            print("Error saving chapter:", e)
            return "The End. (But there was a problem saving the chapter.)"

    if "inventory" in command or "items" in command or "bag" in command:
        inventory = session.get("inventory",
                                ["a revolver", "canteen", "worn journal"])
        return f"You currently have: {', '.join(inventory)}."

    elif "location" in command or "where am i" in command:
        return session.get("current_location",
                           "You're somewhere unfamiliar...")

    elif "goal" in command or "mission" in command:
        return session.get("current_objective",
                           "You're not sure what your goal is yet.")

    elif "clues" in command:
        mystery = session.get("mystery", {})
        clues = mystery.get("clues", [])
        return "Clues so far:\n" + "\n".join(
            f"- {c}"
            for c in clues) if clues else "You haven't found any clues yet."

    elif "suspects" in command:
        mystery = session.get("mystery", {})
        suspects = mystery.get("suspects", [])
        return "Suspects under investigation:\n" + "\n".join(
            f"- {s}"
            for s in suspects) if suspects else "No suspects identified yet."

    elif "deductions" in command or "conclusions" in command:
        mystery = session.get("mystery", {})
        deductions = mystery.get("deductions", [])
        return "Deductions:\n" + "\n".join(
            f"- {d}" for d in deductions
        ) if deductions else "You haven’t made any deductions yet."

    elif command.startswith("accuse"):
        mystery = session.get("mystery", {})
        if not mystery.get("active"):
            return "There's no active mystery to solve."

        if mystery.get("case_status") != "unsolved":
            return "The case is already closed."

        if len(command.split()) < 2:
            return "You must accuse someone. Try: /game accuse Eliora Thorn"

        accused_name = command.split(" ", 1)[1].strip()
        correct = mystery.get("correct_culprit")
        mystery["accusations"].append(accused_name)

        if accused_name.lower() == correct.lower():
            mystery["case_status"] = "solved"
            mystery["active"] = False
            return f"You accuse {accused_name}, and bring them before the court. With the evidence you've gathered, they are found guilty. Justice is served. The case is closed."

        if len(mystery["accusations"]) == 1:
            return f"You accuse {accused_name}, but the case doesn't hold up in court. You have one final chance to accuse the real culprit."

        else:
            mystery["case_status"] = "failed"
            mystery["active"] = False
            return (
                f"You accuse {accused_name} again, but they're acquitted. "
                f"The real culprit — {correct} — walks free. The mystery remains unsolved."
            )

    else:
        return "I didn’t recognize that command. Try asking about your inventory, location, mission, clues, suspects, or deductions."


#For mystery mode, generates dyanmic clues from the story
def generate_dynamic_clue_from_scene(scene_text: str) -> str:
    """
    Uses GPT to generate a subtle, story-embedded clue from a narrative scene.
    The clue should raise suspicion about a suspect without solving the mystery outright.
    """
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role":
            "user",
            "content":
            ("You are an AI writing assistant that creates immersive detective fiction. "
             "Given the following scene from a mystery story, generate a subtle but suspicious clue. "
             "The clue should reference a specific object, phrase, or behavior that might later connect to a suspect, "
             "without revealing the mystery solution. The clue must feel like part of the scene, not exposition.\n\n"
             f"Scene:\n{scene_text}\n\n"
             "Clue (in one immersive sentence):")
        }],
        temperature=0.8,
        max_tokens=80,
    )

    clue = response.choices[0].message.content.strip(
    ) if response.choices[0].message.content else "Failed to generate clue."
    return clue


def update_game_state_from_story(latest_story: str):
    prompt = (
        "You are a game system AI. A new piece of narrative has been generated for a story. "
        "Based only on the most recent passage, extract the following:\n\n"
        "1. The character's current location (brief and specific).\n"
        "2. The current objective or mission, if any.\n\n"
        f"Narrative:\n{latest_story}\n\n"
        "Respond in the following format:\n"
        "Location: <short phrase>\n"
        "Objective: <short phrase or 'Unknown'>")

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.5,
    )

    content = response.choices[0].message.content.strip(
    ) if response.choices[0].message.content else ""
    lines = content.splitlines()
    for line in lines:
        if line.lower().startswith("location:"):
            session["current_location"] = line.split(":", 1)[1].strip()
        elif line.lower().startswith("objective:"):
            session["current_objective"] = line.split(":", 1)[1].strip()


def assess_empathy_score(user_input: str, story: str) -> int:
    prompt = (
        f"On a scale of 1 to 10, where 10 is extremely empathetic and 1 is very selfish, "
        f"rate the following player choice for empathy:\n\n"
        f"Choice: '{user_input}'\n\n"
        f"Recent story context:\n{story}\n\n"
        f"Return ONLY a number from 1 to 10.")

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.3,
    )

    try:
        return int(response.choices[0].message.content.strip() if response.
                   choices[0].message.content is not None else "5")
    except Exception as e:
        print("Error parsing empathy score:", e)
        return 5  # Neutral default if parsing fails


#editor function to improve and customize story generated in /start and /continue
def editor_enhance_story(raw_story: str, genre: str,
                         narrator_style: str) -> str:
    prompt = (
        f"You are an expert literary editor and storyteller. Enhance the following second-person story passage. "
        f"The genre is {genre}, and the narrator should speak in a '{narrator_style}' tone. "
        f"Strengthen immersive language, expand on intriguing details, and improve pacing while preserving original meaning.\n\n"
        f"Story passage:\n{raw_story}\n\n"
        "Return the edited and enhanced version of the story.")

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.7,
    )

    return response.choices[0].message.content.strip() if response.choices[
        0].message.content else "Failed to generate rewritten story."


#create plot outline to give users ability to return to same spot and give AI structure to story
def generate_plot_skeleton(genre: str, character_name: str,
                           story_idea: str) -> list:
    prompt = (
        f"You are a master storyteller. Create a 5–7 step skeleton plot structure for a story. "
        f"The genre is {genre}, and the main character is {character_name}. "
        f"The story idea is: {story_idea or 'Invent your own.'} "
        f"Each step should be a brief 1-2 sentence description of a major plot event or turning point. "
        f"Format as a numbered JSON list like:\n"
        f"[\"Step 1 description\", \"Step 2 description\", ...]")

    response = client.chat.completions.create(
        model="gpt-4o",  # or 3.5 if needed
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.7,
    )

    return json.loads(
        response.choices[0].message.content
    ) if response.choices[0].message.content is not None else []


#to advance the skeleton to the next plot point for saves and story structure
def maybe_advance_plot_step(story_text: str):
    skeleton = session.get("plot_skeleton", [])
    step = session.get("current_step", 0)

    if not skeleton or step >= len(skeleton):
        return  # Nothing to advance

    current_step_text = skeleton[step]
    prompt = (
        f"The following is a story segment and a current plot objective.\n\n"
        f"Objective: {current_step_text}\n\n"
        f"Story Segment:\n{story_text}\n\n"
        f"Has this plot objective been clearly fulfilled in the story? "
        f"Respond with ONLY 'yes' or 'no'.")

    response = client.chat.completions.create(model="gpt-4o",
                                              messages=[{
                                                  "role": "user",
                                                  "content": prompt
                                              }],
                                              temperature=0.3,
                                              max_tokens=10)

    decision = response.choices[0].message.content.strip().lower(
    ) if response.choices[0].message.content is not None else ""
    if decision.startswith("yes"):
        session["current_step"] = step + 1


# Start new story — generate image and return intro + image
@app.post("/start")
def start_story(data: StartRequest):
    session["genre"] = data.genre
    session["character_name"] = data.character_name
    session["character_description"] = data.character_description
    session["history"] = []
    session["story_idea"] = data.story_idea
    session["let_ai_decide"] = data.let_ai_decide
    session["current_location"] = ""
    session["current_objective"] = ""
    session["character_prompt"] = ""
    session["narrator_style"] = data.narrator_style
    session["Empathy"] = [],
    session["Risk-taking"] = [],
    session["Violence tendancy"] = [],
    session["Loyalty"] = [],
    session["Curiosity"] = [],
    session["Rule-breaking"] = [],
    session["Sacrifice"] = [],
    session["Openness"] = [],
    session["Conscientiousness"] = [],
    session["Extroversion"] = [],
    session["Agreeableness"] = [],
    session["Neuroticism"] = [],

    character_prompt = (
        f"Full-body portrait of a character named {data.character_name} from a {data.genre} world. "
        f"{data.character_description}. Dramatic lighting, high detail, cinematic style. No text."
    )
    session["character_prompt"] = character_prompt

    image_url = None
    if not data.skip_image:
        image_url = generate_character_image(character_prompt)
        session["image_url"] = image_url

    story = generate_story(data.genre, character_name=data.character_name)
    #editor function to improve on the story
    #story = editor_enhance_story(story, data.genre, data.narrator_style or "neutral")
    session["history"].append(story)

    # If genre is Mystery, extract suspects and set up investigation
    if "mystery" in data.genre.lower():
        if story is not None:
            suspects = extract_suspects_from_story(story)
            if suspects:
                culprit = random.choice(suspects)
                session["mystery"] = {
                    "title": "Untitled Case",
                    "type": "investigation",
                    "active": True,
                    "suspects": suspects,
                    "clues": [],
                    "false_clues": [],
                    "deductions": [],
                    "case_status": "unsolved",
                    "correct_culprit": culprit,
                    "accusations": []
                }

    session["plot_skeleton"] = generate_plot_skeleton(data.genre,
                                                      data.character_name,
                                                      data.story_idea)
    session["current_step"] = 0  # Start at beginning of the plot skeleton

    update_game_state_from_story(story)
    #Keeps NPC relationship saves updated
    book_title = session.get("book", {}).get("title", "")
    character_name = session.get("character_name", "Unnamed Hero")
    update_npc_memory_from_story(book_title, character_name, story)

    return {"story": story, "image_url": image_url}


# Continue story
@app.post("/continue")
def continue_story(data: ContinueRequest):
    genre = session.get("genre")
    character_name = session.get("character_name")
    book_title = session.get("book", {}).get("title")

    if not genre or not character_name:
        return {"story": "Session not found. Please restart."}

    user_input = data.user_input.strip()

    #Test feature to end the game and create a chapter
    if user_input.lower().strip() == "/game end":
        story = "End of Chapter."
        session["history"].append(story)

        # Create and save the chapter
        try:
            chapter_info = create_chapter_summary_and_story(
                character_name, genre, session["history"])
            save_chapter_to_book(book_title, character_name,
                                 chapter_info["summary"],
                                 chapter_info["prose"])
        except Exception as e:
            print(f"Error generating chapter: {e}")
            return {
                "story":
                "End of Chapter. (But there was a problem saving the chapter.)"
            }

        return {"story": "End of Chapter."}

    if session["history"] and session["history"][-1].strip() == "The End.":
        return {
            "story":
            "The story has already ended. Please return to the bookshelf or start a new chapter."
        }

    if user_input.lower().startswith("/game"):
        command = user_input[5:].strip()
        return {"story": handle_game_command(command)}

    session["history"].append(f"> {user_input}")

    story = generate_story(genre,
                           user_input=user_input,
                           character_name=character_name)
    #editor function to improve on the story
    #story = editor_enhance_story(story, genre, session.get("narrator_style", "neutral"))
    session["history"].append(story)
    #advance skeleton to the next point if we've reached it
    maybe_advance_plot_step(story)

    # Mystery clue generation
    if session.get("mystery", {}).get("active"):
        clue = generate_dynamic_clue_from_scene(story)
        session["mystery"]["clues"].append(clue)

    #Updating saves
    update_game_state_from_story(story)

    #Update NPC relationship info
    character_name = character_name or session.get("character_name",
                                                   "Unnamed Hero")
    book_title = session.get("book", {}).get("title", "")
    update_npc_memory_from_story(book_title, character_name, story)

    # Save the updated session history into the book object
    book = session["book"]
    book["stories"].append({
        "character": character_name,
        "entry": session["history"][-1]
    })
    save_book_data(book)

    character_data = {
        "name": session.get("character_name"),
        "description": session.get("character_description"),
        #"inventory": session.get("inventory", []),
        "reputation": session.get("reputation", {}),
        "relationships": session.get("relationships", {}),
        "history": session.get("history", []),
        "empathy_scores": session.get("empathy_scores", []),
        "character_prompt": session.get("character_prompt", ""),
    }

    #Assess empathy for the current user input
    empathy_score = assess_empathy_score(user_input, story)
    session.setdefault("empathy_scores", []).append(empathy_score)

    #Update character trait data
    traits = analyze_character_traits(user_input, story)

    if traits:
        character_data = load_character(book_title, character_name)
        for trait, score in traits.items():
            # Save to top-level (if still used elsewhere)
            character_data.setdefault(trait, []).append(score)

            # ✅ Also save to trait_scores dict
            if "trait_scores" not in character_data:
                character_data["trait_scores"] = {}
            character_data["trait_scores"].setdefault(trait, []).append(score)
            
            
            #test to see if they're being recorded
            print(f"{trait}: {score}")

    
    #skeleton outline point save in character data
    character_data["current_step"] = session.get("current_step", 0)

    save_character(book_title, character_name, character_data)

    #save plot skeleton progress so users can return to the same spot in the story
    character_data["plot_skeleton"] = session["plot_skeleton"]
    character_data["current_step"] = 0

    #Normal feature to end the story and create a chapter
    if story.strip().endswith("End of Chapter."):
        chapter_info = create_chapter_summary_and_story(
            character_name, genre, session["history"])
        save_chapter_to_book(book_title, character_name,
                             chapter_info["summary"], chapter_info["prose"])

    return {"story": story}


#Save and exit
@app.post("/save-progress")
def save_progress():
    book_title = session.get("book", {}).get("title")
    character_name = session.get("character_name")
    if not book_title or not character_name:
        return {"status": "error", "message": "No story in progress."}

    character_data = {
        "name": character_name,
        "description": session.get("character_description", ""),
        #"inventory": session.get("inventory", []),
        "reputation": session.get("reputation", {}),
        "relationships": session.get("relationships", {}),
        "empathy_scores": session.get("empathy_scores", []),
        "plot_skeleton": session.get("plot_skeleton", []),
        "current_step": session.get("current_step", 0),
    }

    #skeleton outline save
    character_data["current_step"] = session.get("current_step", 0)

    save_character(book_title, character_name, character_data)

    return {"status": "ok", "message": "Progress saved."}


@app.post("/generate-moment-image")
def generate_moment_image():
    genre = session.get("genre")
    character_name = session.get("character_name")
    character_prompt = session.get("character_prompt", "")
    history = session.get("history", [])

    # Use last 2 story blocks as moment context
    moment_context = "\n".join(history[-2:])

    # Clean, non-redundant image prompt using saved character_prompt
    image_prompt = (
        f"Illustration of a dramatic moment from a {genre} story. "
        f"The character is {character_name}. Preserve their appearance exactly as described here: {character_prompt} "
        f"Scene: {moment_context}. Make it cinematic, vivid, and emotional. No text or logos."
    )

    image_url = generate_character_image(image_prompt)
    return {"image_url": image_url}


# --- FIXED LOAD_BOOK ROUTE ---
@app.post("/load-book")
def load_book_route(data: BookLoadRequest):
    book = load_book_data(data.title)
    if not book:
        book = {
            "title": data.title,
            "genre": data.genre,
            "characters": [],
            "stories": [],
        }

    # Fix: Ensure characters is a list
    if "characters" not in book or not isinstance(book["characters"], list):
        book["characters"] = []

    if data.character_name not in book["characters"]:
        book["characters"].append(data.character_name)

    # Also ensure stories is a list
    if "stories" not in book or not isinstance(book["stories"], list):
        book["stories"] = []

    save_book_data(book)

    # Load or create character and their inventory

    character_data = create_or_load_character(
        book_title=data.title,
        character_name=data.character_name,
        genre=book["genre"],
        description=data.character_description or "")
    # Store session
    session["genre"] = book["genre"]
    session["character_name"] = data.character_name
    session["character_description"] = data.character_description
    session["book"] = book
    session["history"] = character_data.get("history", [])
    #session["inventory"] = character_data.get("inventory", [])
    session["reputation"] = character_data.get("reputation", {})
    session["relationships"] = character_data.get("relationships", {})
    session["story_idea"] = data.story_idea or ""
    session["let_ai_decide"] = data.let_ai_decide
    session["character"] = character_data
    session["empathy_scores"] = character_data.get("empathy_scores", [])
    session["plot_skeleton"] = character_data.get("plot_skeleton", [])
    session["current_step"] = character_data.get("current_step", 0)
    session["current_step"] = character_data.get("current_step", 0)

    image_url = None
    if not data.skip_image:
        character_prompt = (
            f"Full-body portrait of a character named {data.character_name} from a {book['genre']} world. "
            f"{data.character_description}. Dramatic lighting, high detail, cinematic style. No text."
        )
        session["character_prompt"] = character_prompt
        image_url = generate_character_image(character_prompt)
        session["image_url"] = image_url

    # Begin story
    story = generate_story(genre=book["genre"],
                           character_name=data.character_name)
    session["history"].append(story)
    update_game_state_from_story(story)

    return {"story": story, "image_url": image_url}


@app.get("/list-books")
def list_books():
    books = []
    for filename in os.listdir(BOOKS_DIR):
        if filename.endswith(".json"):
            path = os.path.join(BOOKS_DIR, filename)
            with open(path, "r") as f:
                data = json.load(f)
                books.append({
                    "title": data.get("title", "Unknown"),
                    "genre": data.get("genre", "Unknown"),
                    "characters": data.get("characters", []),
                    "chapters": data.get("chapters", [])
                })
    return {"books": books}


@app.post("/export-story")
def export_story():
    # Gather session variables
    genre = session.get("genre", "Unknown")
    character_name = session.get("character_name", "Unnamed Hero")
    character_description = session.get("character_description", "")
    history = session.get("history", [])

    # Step 1: Filter out /game commands and responses
    # This step messes up the Word export, using the full session history in the prompt instead of filtered history.
    #filtered_history = []
    #skip_next = False
    #for line in history:
        #if line.strip().lower().startswith("/game") or skip_next:
            #skip_next = not skip_next  # Skip the AI's system response too
            #continue
        #filtered_history.append(line)

    #full_log = "\n".join(filtered_history)

    # Step 2: Ask GPT to rewrite it into 3rd-person past tense
    prompt = (
        f"Rewrite the following story log as a third-person limited past-tense narrative. "
        f"The protagonist is named {character_name}. Genre: {genre}. "
        f"Only include what {character_name} could reasonably know or perceive. "
        f"Describe other characters’ actions and dialogue in vivid, cinematic prose. "
        f"Remove all second-person references, user prompts, and transform it into polished storybook prose. "
        f"Here is the story log:\n\n{history}")

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.7,
    )

    rewritten_story = response.choices[0].message.content.strip(
    ) if response.choices[0].message.content is not None else ""

    # Step 3: Create the Word document
    doc = Document()

    # Title Page
    doc.add_heading(f"{character_name}'s Story", level=0)
    doc.add_paragraph(f"Genre: {genre}")
    doc.add_paragraph(f"Character Description: {character_description}")
    doc.add_paragraph()  # blank line

    # Story Content
    paragraphs = rewritten_story.split("\n\n")
    for para in paragraphs:
        paragraph = doc.add_paragraph()  # Don't style the paragraph directly
        run = paragraph.add_run(para.strip())
        font = run.font
        font.name = 'Georgia'
        font.size = Pt(12)

    # Save file
    filename = f"story_export_{uuid.uuid4().hex}.docx"
    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    doc.save(filepath)

    return FileResponse(
        filepath,
        media_type=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename)


@app.get("/get-chapters/{book_title}")
def get_chapters(book_title: str):
    book = load_book_data(book_title)
    if not book or "chapters" not in book:
        return {"chapters": []}

    chapters = []
    for idx, chapter in enumerate(book["chapters"]):
        chapters.append({
            "number":
            idx + 1,
            "character":
            chapter.get("character", "Unknown"),
            "summary":
            chapter.get("summary", "No summary available"),
        })

    return {"chapters": chapters}


@app.get("/export-chapter/{book_title}/{chapter_index}")
def export_chapter(book_title: str, chapter_index: int):
    book = load_book_data(book_title)
    if not book or "chapters" not in book or chapter_index >= len(book["chapters"]):
        raise HTTPException(status_code=404, detail="Chapter not found")

    chapter = book["chapters"][chapter_index]
    character_name = chapter.get("character", "Unnamed Hero")
    genre = book.get("genre", "Unknown")
    story_text = chapter.get("story", "")

    # ✅ REWRITE the raw story into third-person narrative
    prompt = (
        f"Rewrite the following story as a third-person limited past-tense narrative. "
        f"The protagonist is named {character_name}. Genre: {genre}. "
        f"Only include what {character_name} could reasonably know or perceive. "
        f"Describe other characters’ actions and dialogue in vivid, cinematic prose. "
        f"Remove second-person references and reframe as polished storybook prose.\n\n"
        f"{story_text}"
    )

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        rewritten_story = response.choices[0].message.content.strip() if response.choices[0].message.content else ""
    except Exception as e:
        print("GPT rewrite failed:", e)
        rewritten_story = story_text  # fallback

    doc = Document()
    doc.add_heading(f"{character_name}'s Chapter {chapter_index + 1}", level=0)
    doc.add_paragraph(f"Genre: {genre}")
    doc.add_paragraph()  # spacer

    for para in rewritten_story.split("\n\n"):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(para.strip())
        run.font.size = Pt(12)
        run.font.name = 'Georgia'

    # --- ✅ Improved filename ---
    safe_title = book_title.replace(" ", "_")
    safe_character = character_name.replace(" ", "_")
    filename = f"{safe_title}_{safe_character}_Chapter_{chapter_index + 1}.docx"

    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    doc.save(filepath)

    with open(filepath, "rb") as f:
        file_content = f.read()
        encoded_file = base64.b64encode(file_content).decode("utf-8")

    return {
        "filename": filename,
        "filedata": encoded_file,
    }


@app.post("/test-chapter")
def test_chapter():
    genre = session.get("genre")
    character_name = session.get("character_name")
    book = session.get("book")
    history = session.get("history", [])

    if not genre or not character_name or not book:
        return {
            "status":
            "error",
            "message":
            f"Missing session data: genre={genre}, character_name={character_name}, book={book}"
        }

    book_title = book.get("title")
    if not book_title:
        return {
            "status": "error",
            "message": "Book title is missing from session['book']."
        }

    try:
        chapter_info = create_chapter_summary_and_story(
            character_name, genre, history)

        save_chapter_to_book(book_title, character_name,
                             chapter_info["summary"], chapter_info["prose"])

        return {
            "status": "success",
            "message": "Chapter created and saved successfully!"
        }

    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to create chapter: {str(e)}"
        }


@app.get("/developer/empathy-scores")
def get_empathy_scores():
    results = []
    total_scores = []

    for filename in os.listdir(CHARACTERS_DIR):
        if filename.endswith(".json"):
            with open(os.path.join(CHARACTERS_DIR, filename), "r") as f:
                char_data = json.load(f)
                empathy_scores = char_data.get("empathy_scores", [])
                if empathy_scores:
                    average = round(
                        sum(empathy_scores) / len(empathy_scores), 2)
                    results.append({
                        "character":
                        char_data.get("name", "Unknown"),
                        "book":
                        filename.split("__")[0].replace("_", " "),
                        "average":
                        average,
                        "count":
                        len(empathy_scores),
                        "all_scores":
                        empathy_scores,
                    })
                    total_scores.extend(empathy_scores)

    overall_avg = round(sum(total_scores) /
                        len(total_scores), 2) if total_scores else 0
    return {"characters": results, "overall_average": overall_avg}


@app.get("/developer/trait-scores")
def get_trait_scores():
    import os
    import json

    TRAITS = [
        "empathy", "risk-taking", "violence-tendency", "loyalty", "curiosity",
        "rule-breaking", "sacrifice", "openness", "conscientiousness",
        "extroversion", "agreeableness", "neuroticism"
    ]

    all_trait_totals = {trait: [] for trait in TRAITS}
    character_results = []

    for filename in os.listdir(CHARACTERS_DIR):
        if filename.endswith(".json"):
            with open(os.path.join(CHARACTERS_DIR, filename), "r") as f:
                char_data = json.load(f)

                # Load all trait score histories from character data
                trait_scores = char_data.get("trait_scores", {})
                char_result = {
                    "character": char_data.get("name", "Unknown"),
                    "book": filename.split("__")[0].replace("_", " "),
                    "traits": {}
                }

                for trait in TRAITS:
                    scores = trait_scores.get(trait, [])
                    if scores:
                        avg = round(sum(scores) / len(scores), 2)
                        char_result["traits"][trait] = {
                            "average": avg,
                            "count": len(scores),
                            "all_scores": scores
                        }
                        all_trait_totals[trait].extend(scores)

                character_results.append(char_result)

    # Calculate overall averages
    overall_averages = {
        trait: round(sum(scores) / len(scores), 2) if scores else 0
        for trait, scores in all_trait_totals.items()
    }

    return {
        "characters": character_results,
        "overall_averages": overall_averages
    }


@app.post("/start-mystery")
def start_mystery(case_title: Optional[str] = Body(None),
                  case_type: Optional[str] = Body(None)):
    book_title = session.get("book", {}).get("title")
    character_name = session.get("character_name")
    genre = "Mystery"

    if not book_title or not character_name:
        raise HTTPException(status_code=400,
                            detail="No active book or character.")

    suspects = generate_suspects(genre)
    first_clue = generate_initial_clue(suspects)

    mystery_data = {
        "title": case_title or "Untitled Mystery",
        "type": case_type or "Unknown Type",
        "suspects": suspects,
        "clues": [first_clue],
        "culprit": None,
        "solved": False
    }

    save_mystery_state(book_title, character_name, mystery_data)
    session["mystery"] = mystery_data  # ← Also save in session if needed

    return {
        "status": "mystery_started",
        "title": mystery_data["title"],
        "type": mystery_data["type"],
        "first_clue": first_clue,
        "suspects": suspects
    }


# Call this with your OpenAI key loaded elsewhere


def analyze_character_traits(user_input: str, story_context: str) -> dict:
    prompt = f"""
You are a narrative psychologist. Based on the following player's choice and the story context, rate the player's traits from 1 to 10, where 1 means very low and 10 means very high expression of the trait.

Traits to score:
- Empathy (how compassionate and selfless the action is)
- Risk-taking (willingness to take bold or dangerous actions)
- Violence tendency (use or threat of violence)
- Loyalty (commitment to allies or principles)
- Curiosity (desire to explore, learn, or investigate)
- Rule-breaking (disregard for authority or social norms)
- Sacrifice (willingness to endure loss for others)
- Openness (creativity, new ideas, emotional openness)
- Conscientiousness (carefulness, responsibility)
- Extroversion (initiative, sociability)
- Agreeableness (kindness, cooperativeness)
- Neuroticism (emotional instability, impulsiveness)

Provide ONLY the scores in valid JSON format like this:
{{"empathy": 8, "risk-taking": 6, "violence-tendency": 3, "loyalty": 9, "curiosity": 7, "rule-breaking": 4, "sacrifice": 6, "openness": 5, "conscientiousness": 7, "extroversion": 4, "agreeableness": 8, "neuroticism": 3}}

Player's Choice: "{user_input}"

Story Context:
{story_context}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{
                "role": "system",
                "content": "You are an expert character trait analyzer."
            }, {
                "role": "user",
                "content": prompt
            }],
            temperature=0.3,
        )

        text = response.choices[0].message.content.strip(
        ) if response.choices[0].message.content is not None else ""
        traits = json.loads(text)
        return traits

    except Exception as e:
        print("Failed to parse traits:", e)
        return {}


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
